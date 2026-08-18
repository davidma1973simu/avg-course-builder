# -*- coding: utf-8 -*-
"""
生成《核心设计假设（锁定）》文档：
  docs/Core_Design_Assumption.html + .docx
单源，避免漂移。
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(__file__), "..", "docs")
os.makedirs(OUT, exist_ok=True)

TITLE = "核心设计假设（锁定）"
SUB = "AVG 探索世界默认采用「远」模式（虚构类比情景）+ 强制「应用迁移桥」"

SECTIONS = [
    ("1 · 一句话假设", [
        "AVG 的探索世界**默认采用「远」（虚构类比）情景**，作为心理安全的练习场；学员在虚构世界中体验、"
        "探索、试错，沉淀出方法论与工具流程，再**强制迁移应用到学员自身的真实工作场景**。",
        "无论采用真实还是虚构场景，复盘（Reflection）之后都必须包含一个「将所学认知 / 方法 / 技能 / 流程 / "
        "工具应用到真实场景与挑战」的环节——**没有这条，学习与产品不成立。**",
    ]),
    ("2 · 近 vs 远：定义与平台选择", [
        "**近（仿真真实）**：场景就是学员的工作现场（会议室、客户拜访、绩效面谈）。\n"
        "优点：直接、易共鸣。缺点：ego 威胁大（怕暴露无能）、创新受限、难以抽象出通用方法。",
        "**远（虚构类比）**：场景是类比世界（星际探索、古堡谜案、末日避难所），用「距离」制造心理安全，"
        "学员更敢冒险、更易把具体经历抽象成通用方法。\n"
        "缺点：需要做「翻译」——把虚构体验连回真实工作。",
        "**平台默认 = 远**，并强制「应用迁移桥」补上翻译环节。客户如有强合规 / 强真实诉求可选「近」，"
        "但**应用迁移桥不可省**。",
    ]),
    ("3 · 应用迁移桥（强制环节 · Mandatory）", [
        "位置：Reflection 阶段的最后一个子步，或独立成「应用（Application）」阶段。",
        "内容：学员带着在虚构世界中沉淀的方法论 / 工具 / 流程，回到**自己真实的一个工作挑战**，完成：",
        "· 我的真实挑战是什么（来自客户定制场景 / 学员自身）  \n"
        "· 我会用哪个方法 / 工具  \n"
        "· 第一步行动是什么（有时限）  \n"
        "· 产出「真实场景应用卡 / 迁移行动计划」（离线包可带走产出物之一）",
    ]),
    ("4 · 对工作流 / Pipeline / 离线包的影响", [
        "**Pipeline**：在 Reflection 后（或其中）明确「应用迁移」子步；Ending 的 abilityRadar 要能映射到真实能力维度。",
        "**离线包**：必须含「真实场景应用卡」模板（由 VR-RE05 校验）；DM 手册含「从虚构到真实」的引导话术。",
        "**校验**：新增 VR-RE05（BLOCKER）——缺应用迁移桥，产品不得发布。",
    ]),
    ("5 · 示例：团队决策与冲突管理 → 星际探索", [
        "课程主题：团队决策与冲突管理。",
        "我们选择：**星际探索场景**（飞船登陆未知星球，机组分歧、资源冲突、信任崩塌），而非工作场景。",
        "用途：在星际场景中体验和探索「分歧如何升级为冲突、不同决策风格的后果」，产出「冲突诊断—干预—共识」的"
        "方法论与工具流程；然后**应用迁移桥**：学员把这套方法带回到自己团队的真实决策冲突中，写出第一步行动。",
    ]),
    ("6 · 校验规则（锁定）", [
        "**VR-RE05（复盘 / 应用，BLOCKER）**：stages.reflection 或 stages.application 必须含 `realScenarioApplication` "
        "字段（真实挑战 + 选用方法 + 有时限第一步），且离线包 `kit.takeawayTemplates` 含 `'realScenarioApplication'` 模板。",
        "失败信息：缺少「应用迁移桥」，所学无法落地到真实工作。",
    ]),
]

CSS = """
<style>
:root{--green:#0E6B50;--ink:#1f2933;--mut:#5b6770;--line:#e3e8ec;--bg:#f7f9fa;}
*{box-sizing:border-box;}
body{font-family:-apple-system,"Microsoft YaHei","PingFang SC",Segoe UI,Roboto,sans-serif;
 color:var(--ink);line-height:1.75;margin:0;background:var(--bg);}
.wrap{max-width:920px;margin:0 auto;padding:48px 28px 80px;background:#fff;
 box-shadow:0 1px 3px rgba(0,0,0,.06);min-height:100vh;}
h1{color:var(--green);font-size:30px;margin:0 0 6px;letter-spacing:.5px;}
.sub{color:var(--mut);font-size:15px;margin:0 0 28px;font-weight:500;}
h2{color:var(--green);font-size:20px;margin:34px 0 10px;padding-bottom:6px;border-bottom:2px solid var(--line);}
p{margin:10px 0;font-size:15px;}
strong{color:#0b3d2e;}
.toc{background:var(--bg);border:1px solid var(--line);border-radius:10px;padding:16px 22px;margin:0 0 30px;}
.toc b{color:var(--green);}
.toc ol{margin:8px 0 0;padding-left:22px;}
.toc li{margin:4px 0;font-size:14px;}
footer{margin-top:48px;color:var(--mut);font-size:12.5px;border-top:1px solid var(--line);padding-top:14px;}
</style>
"""

def esc(s):
    return s.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

def md_inline(s):
    # support **bold** and \n
    s = esc(s)
    import re
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    s = s.replace("\n", "<br>")
    return s

def build_html(path):
    toc = "".join(f"<li>{esc(t)}</li>" for t,_ in SECTIONS)
    body = ""
    for t, paras in SECTIONS:
        body += f"<h2>{esc(t)}</h2>\n"
        for p in paras:
            body += f"<p>{md_inline(p)}</p>\n"
    html = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{esc(TITLE)}</title>{CSS}</head><body><div class="wrap">
<h1>{esc(TITLE)}</h1><p class="sub">{esc(SUB)}</p>
<div class="toc"><b>目录</b><ol>{toc}</ol></div>
{body}
<footer>AVG Course Builder · 核心设计假设（锁定） · 本假设为产品不可动摇的前提，缺应用迁移桥则学习与产品不成立。</footer>
</div></body></html>"""
    with open(path,"w",encoding="utf-8") as f:
        f.write(html)

def set_cjk(run, name="Microsoft YaHei"):
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def build_docx(path):
    doc = Document()
    st = doc.styles['Normal']; st.font.size = Pt(11)
    t = doc.add_heading(TITLE, level=0)
    for r in t.runs: set_cjk(r)
    sp = doc.add_paragraph(SUB)
    for r in sp.runs: set_cjk(r); r.font.size = Pt(11); r.font.italic = True; r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    for t_, paras in SECTIONS:
        h = doc.add_heading(t_, level=1)
        for r in h.runs: set_cjk(r)
        for p in paras:
            # split on \n for line breaks
            pp = doc.add_paragraph()
            parts = p.split("\n")
            for i, seg in enumerate(parts):
                if i>0:
                    pp.add_run().add_break()
                # parse **bold** inline
                import re as _re
                toks = _re.split(r"(\*\*.*?\*\*)", seg)
                for tok in toks:
                    if not tok:
                        continue
                    if tok.startswith("**") and tok.endswith("**"):
                        run = pp.add_run(tok[2:-2]); run.bold = True
                    else:
                        run = pp.add_run(tok)
                    set_cjk(run)
    doc.save(path)

if __name__ == "__main__":
    h = os.path.abspath(os.path.join(OUT,"Core_Design_Assumption.html"))
    d = os.path.abspath(os.path.join(OUT,"Core_Design_Assumption.docx"))
    build_html(h)
    build_docx(d)
    print("HTML ->", h)
    print("DOCX ->", d)
