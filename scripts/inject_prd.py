# -*- coding: utf-8 -*-
"""向 PRD.html / PRD.docx 注入「核心设计假设（锁定）」加注段（源 content_prd.py 已回收，仅做追加，不重写）。"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

B = "/Users/davidma/WorkBuddy/2026-08-16-15-33-43/AVG-Course-Builder/docs"
HTML = os.path.join(B, "PRD.html")
DOCX = os.path.join(B, "PRD.docx")

HTML_BLOCK = """
<section id="core-assumption" style="margin-top:40px;border:2px solid #0E6B50;border-radius:10px;padding:18px 22px;background:#f3faf7;">
<h2 style="color:#0E6B50;margin-top:0;">核心设计假设（锁定 · 不可动摇）</h2>
<p><strong>一句话：</strong>AVG 的探索世界默认采用「远」（虚构类比）情景，作为心理安全的练习场；学员在虚构中体验 / 试错、沉淀方法论与工具，
复盘后<strong>必须强制「应用迁移桥」</strong>——把所学迁移到学员自身真实工作场景。无论近 / 远，<strong>没有这条，学习与产品不成立。</strong></p>
<ul>
  <li><strong>远（默认）</strong>：星际探索 / 古堡谜案等类比世界，用距离降低 ego 威胁、提升试错与抽象能力。</li>
  <li><strong>近（可选）</strong>：仿真真实工作现场；仅在客户强合规 / 强真实诉求时选用，但<strong>应用迁移桥不可省</strong>。</li>
  <li><strong>应用迁移桥（强制）</strong>：真实挑战 + 选用方法 + 有时限第一步 → 离线包内嵌「真实场景应用卡」。</li>
  <li><strong>校验</strong>：新增 <code>VR-RE05</code>（BLOCKER），缺桥不得发布。详见 <a href="Core_Design_Assumption.html">Core_Design_Assumption</a>。</li>
</ul>
</section>
"""

def inject_html():
    s = open(HTML, encoding="utf-8").read()
    if "核心设计假设（锁定" in s:
        print("PRD.html 已含核心假设段，跳过")
        return
    s = s.replace("</body>", HTML_BLOCK + "\n</body>", 1)
    open(HTML, "w", encoding="utf-8").write(s)
    print("PRD.html 已注入核心假设段")

def set_cjk(run, name="Microsoft YaHei"):
    run.font.name = name
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {}); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def inject_docx():
    d = Document(DOCX)
    # skip if already present
    if any("核心设计假设（锁定" in p.text for p in d.paragraphs):
        print("PRD.docx 已含核心假设段，跳过"); return
    h = d.add_heading("核心设计假设（锁定 · 不可动摇）", level=1)
    for r in h.runs: set_cjk(r)
    p1 = d.add_paragraph("AVG 的探索世界默认采用「远」（虚构类比）情景，作为心理安全的练习场；学员在虚构中体验/试错、沉淀方法论与工具，"
                         "复盘后必须强制「应用迁移桥」——把所学迁移到学员自身真实工作场景。无论近/远，没有这条，学习与产品不成立。")
    for r in p1.runs: set_cjk(r)
    for t in ["远（默认）：星际探索/古堡谜案等类比世界，用距离降低 ego 威胁、提升试错与抽象能力。",
              "近（可选）：仿真真实工作现场；仅在客户强合规/强真实诉求时选用，但应用迁移桥不可省。",
              "应用迁移桥（强制）：真实挑战 + 选用方法 + 有时限第一步 → 离线包内嵌「真实场景应用卡」。",
              "校验：新增 VR-RE05（BLOCKER），缺桥不得发布。详见 Core_Design_Assumption.docx。"]:
        p = d.add_paragraph(t, style='List Bullet')
        for r in p.runs: set_cjk(r)
    d.save(DOCX)
    print("PRD.docx 已注入核心假设段")

if __name__ == "__main__":
    inject_html()
    inject_docx()
