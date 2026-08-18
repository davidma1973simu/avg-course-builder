# -*- coding: utf-8 -*-
"""
从 Skill 的 references/ 源（Markdown）渲染 Capability_Assets.html + .docx。
单源：Skill md 是事实源，此脚本只做呈现。重建版（原 render_assets.py 已被环境回收）。
"""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = os.path.join(os.path.dirname(__file__), "..")
SKILL = os.path.join(ROOT, "skill", "avg-scenario-design-kit")
OUT = os.path.join(ROOT, "docs")
os.makedirs(OUT, exist_ok=True)

SECTIONS = [
    ("SKILL.md", "00 · 总览与七条心法"),
    ("references/avg_mechanics_library.md", "01 · AVG 互动与设计机制库"),
    ("references/jubensha_playbook.md", "02 · 剧本杀企业培训方法论"),
    ("references/case_library.md", "03 · 案例库"),
    ("references/design_checklists.md", "04 · 设计期检查清单"),
    ("references/validation_rules.md", "05 · 自动校验规则"),
    ("references/offline_package_spec.md", "06 · 线下产品包实体物料规格"),
]

GREEN = "0E6B50"

def inline(s):
    s = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    s = re.sub(r"`([^`]+)`", r"<code>\1</code>", s)
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    return s

def parse_md(text):
    """Return list of blocks: ('h1'|'h2'|'h3'|'p'|'ul'|'ol'|'quote'|'hr'|'table', data)."""
    lines = text.split("\n")
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if line.strip() == "":
            i += 1; continue
        if line.strip() == "---":
            blocks.append(("hr", None)); i += 1; continue
        if line.startswith("```"):
            i += 1; buf = []
            while i < n and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            blocks.append(("code", "\n".join(buf))); continue
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip())); i += 1; continue
        if line.startswith("## "):
            blocks.append(("h2", line[3:].strip())); i += 1; continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip())); i += 1; continue
        if line.startswith("> "):
            buf = []
            while i < n and lines[i].startswith(">"):
                buf.append(lines[i][2:].strip()); i += 1
            blocks.append(("quote", " ".join(buf))); continue
        if line.startswith("|"):
            rows = []
            while i < n and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
            if len(rows) >= 2 and re.match(r"^[\s:|-]+$", rows[1][0] if rows[1] else ""):
                rows = [rows[0]] + rows[2:]
            blocks.append(("table", rows)); continue
        if re.match(r"^\s*[-*] ", line):
            buf = []
            while i < n and re.match(r"^\s*[-*] ", lines[i]):
                buf.append(re.sub(r"^\s*[-*] ", "", lines[i])); i += 1
            blocks.append(("ul", buf)); continue
        if re.match(r"^\s*\d+\. ", line):
            buf = []
            while i < n and re.match(r"^\s*\d+\. ", lines[i]):
                buf.append(re.sub(r"^\s*\d+\. ", "", lines[i])); i += 1
            blocks.append(("ol", buf)); continue
        # paragraph (collect consecutive non-special lines)
        buf = [line]
        i += 1
        while i < n and lines[i].strip() != "" and not lines[i].startswith(("#","|",">","-","*","`"," ")) and not re.match(r"^\s*\d+\. ", lines[i]) and lines[i].strip() != "---":
            buf.append(lines[i]); i += 1
        blocks.append(("p", " ".join(buf)))
    return blocks

# ---------- HTML ----------
def build_html(path):
    toc = ""
    sections_html = ""
    for idx, (fname, title) in enumerate(SECTIONS):
        fp = os.path.join(SKILL, fname)
        text = open(fp, encoding="utf-8").read()
        blocks = parse_md(text)
        # TOC from h2 of this section
        toc += f'<li><a href="#sec{idx}">{title}</a></li>'
        sec = f'<section id="sec{idx}"><h2 class="sech">{title}</h2>'
        for kind, data in blocks:
            if kind == "h1":
                sec += f'<h1>{inline(data)}</h1>'
            elif kind == "h2":
                sec += f'<h3 class="subh">{inline(data)}</h3>'
            elif kind == "h3":
                sec += f'<h4>{inline(data)}</h4>'
            elif kind == "p":
                sec += f'<p>{inline(data)}</p>'
            elif kind == "quote":
                sec += f'<blockquote>{inline(data)}</blockquote>'
            elif kind == "hr":
                sec += "<hr>"
            elif kind == "code":
                sec += f'<pre>{inline(data)}</pre>'
            elif kind == "ul":
                sec += "<ul>" + "".join(f"<li>{inline(x)}</li>" for x in data) + "</ul>"
            elif kind == "ol":
                sec += "<ol>" + "".join(f"<li>{inline(x)}</li>" for x in data) + "</ol>"
            elif kind == "table":
                sec += "<table><thead><tr>" + "".join(f"<th>{inline(c)}</th>" for c in data[0]) + "</tr></thead><tbody>"
                for row in data[1:]:
                    sec += "<tr>" + "".join(f"<td>{inline(c)}</td>" for c in row) + "</tr>"
                sec += "</tbody></table>"
        sec += "</section>"
        sections_html += sec
    css = f"""
<style>
:root{{--green:#{GREEN};--ink:#1f2933;--mut:#5b6770;--line:#e3e8ec;--bg:#f7f9fa;}}
*{{box-sizing:border-box;}}
body{{font-family:-apple-system,"Microsoft YaHei","PingFang SC",Segoe UI,Roboto,sans-serif;color:var(--ink);line-height:1.7;margin:0;background:var(--bg);}}
.wrap{{max-width:960px;margin:0 auto;padding:48px 28px 80px;background:#fff;box-shadow:0 1px 3px rgba(0,0,0,.06);min-height:100vh;}}
h1{{color:var(--green);font-size:26px;margin:18px 0 8px;}}
.sech{{color:var(--green);font-size:22px;margin:40px 0 10px;padding-bottom:6px;border-bottom:2px solid var(--line);}}
.subh{{color:#0b3d2e;font-size:17px;margin:22px 0 6px;}}
h4{{font-size:14.5px;margin:16px 0 4px;color:var(--ink);}}
p{{margin:9px 0;font-size:14.5px;}}
ul,ol{{margin:9px 0;padding-left:24px;font-size:14.5px;}}
li{{margin:4px 0;}}
blockquote{{background:var(--bg);border-left:4px solid var(--green);margin:12px 0;padding:10px 16px;color:var(--mut);border-radius:0 6px 6px 0;}}
code{{background:#eef2f4;padding:1px 5px;border-radius:4px;font-size:13px;}}
pre{{background:#0f1b22;color:#e6edf3;padding:14px 16px;border-radius:8px;overflow:auto;font-size:12.5px;}}
table{{border-collapse:collapse;width:100%;margin:14px 0;font-size:13.5px;}}
th,td{{border:1px solid var(--line);padding:7px 10px;text-align:left;vertical-align:top;}}
th{{background:var(--bg);color:var(--green);font-weight:700;}}
tr:nth-child(even) td{{background:#fafcfd;}}
.toc{{background:var(--bg);border:1px solid var(--line);border-radius:10px;padding:16px 22px;margin:0 0 30px;}}
.toc b{{color:var(--green);}}
.toc ol{{margin:8px 0 0;}}
footer{{margin-top:48px;color:var(--mut);font-size:12.5px;border-top:1px solid var(--line);padding-top:14px;}}
</style>"""
    html = f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1"><title>AVG × 剧本杀 设计能力资产库</title>{css}</head>
<body><div class="wrap"><h1>AVG × 剧本杀 场景化学习设计能力资产库</h1>
<p class="sub" style="color:var(--mut);font-size:14px;">可调用 Skill 的呈现版（源：skill/avg-scenario-design-kit）。含核心设计假设、AVG 机制、剧本杀方法论、案例、检查清单、校验规则、物料规格。</p>
<div class="toc"><b>目录</b><ol>{toc}</ol></div>
{sections_html}
<footer>AVG Course Builder · 设计能力资产库（单源来自 Skill，重渲染即同步）。核心假设：远场景 + 应用迁移桥，缺桥则学习与产品不成立。</footer>
</div></body></html>"""
    open(path, "w", encoding="utf-8").write(html)

# ---------- DOCX ----------
def set_cjk(run, name="Microsoft YaHei"):
    run.font.name = name
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {}); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def build_docx(path):
    doc = Document()
    st = doc.styles['Normal']; st.font.size = Pt(10.5)
    doc.add_heading("AVG × 剧本杀 场景化学习设计能力资产库", level=0)
    for fname, title in SECTIONS:
        fp = os.path.join(SKILL, fname)
        text = open(fp, encoding="utf-8").read()
        blocks = parse_md(text)
        h = doc.add_heading(title, level=1)
        for r in h.runs: set_cjk(r)
        for kind, data in blocks:
            if kind == "h1":
                h = doc.add_heading(data, level=1); 
                for r in h.runs: set_cjk(r)
            elif kind == "h2":
                h = doc.add_heading(data, level=2)
                for r in h.runs: set_cjk(r)
            elif kind == "h3":
                h = doc.add_heading(data, level=3)
                for r in h.runs: set_cjk(r)
            elif kind == "p":
                p = doc.add_paragraph(data)
                for r in p.runs: set_cjk(r)
            elif kind == "quote":
                p = doc.add_paragraph(data); p.style = doc.styles['Intense Quote']
                for r in p.runs: set_cjk(r)
            elif kind == "hr":
                doc.add_paragraph("─" * 20)
            elif kind == "code":
                p = doc.add_paragraph(data); 
                for r in p.runs: r.font.name = "Consolas"; r.font.size = Pt(9)
            elif kind == "ul":
                for x in data:
                    p = doc.add_paragraph(x, style='List Bullet')
                    for r in p.runs: set_cjk(r)
            elif kind == "ol":
                for x in data:
                    p = doc.add_paragraph(x, style='List Number')
                    for r in p.runs: set_cjk(r)
            elif kind == "table":
                t = doc.add_table(rows=len(data), cols=len(data[0])); t.style = 'Table Grid'
                for ri, row in enumerate(data):
                    for ci, cell in enumerate(row):
                        t.cell(ri, ci).text = cell
    doc.save(path)

if __name__ == "__main__":
    h = os.path.abspath(os.path.join(OUT, "Capability_Assets.html"))
    d = os.path.abspath(os.path.join(OUT, "Capability_Assets.docx"))
    build_html(h)
    build_docx(d)
    print("HTML ->", h)
    print("DOCX ->", d)
