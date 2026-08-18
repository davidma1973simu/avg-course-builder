# -*- coding: utf-8 -*-
"""向 PRD.html / PRD.docx 注入「课程分层 3×2 矩阵（设计要求 · 设计前前置判定）」章节。"""
import os
from docx import Document
from docx.oxml.ns import qn

B = "/Users/davidma/WorkBuddy/2026-08-16-15-33-43/AVG-Course-Builder/docs"
HTML = os.path.join(B, "PRD.html")
DOCX = os.path.join(B, "PRD.docx")

HTML_BLOCK = """
<section id="course-tiering" style="margin-top:40px;border:2px solid #1e3a5f;border-radius:10px;padding:18px 22px;background:#f0f4f9;">
<h2 style="color:#1e3a5f;margin-top:0;">课程分层 3×2 矩阵（设计要求 · 设计前前置判定）</h2>
<p><strong>设计要求：</strong>任何课程进入平台，<strong>必须先完成【课程准入与分层判定】（Course Intake &amp; Tiering）</strong>，
输出 tier / track 判定后，主流程（Pipeline）才启动。该判定是 Pipeline 的<strong>第 0 阶段前置输入</strong>，不是可选步骤。
完整矩阵与盲区见 <a href="Course_Tiered_Matrix.html">Course_Tiered_Matrix</a>。</p>
<ul>
  <li><strong>Tier（3 层）</strong>：L1 轻量切片 3h / L2 标准单元 7h / L3 深度系列 14h；<strong>AVG 探索体验严格占总时长 50%</strong>（<code>VR-TIER01</code> BLOCKER）。</li>
  <li><strong>Track（2 类）</strong>：C1 通用改编（已有课程转译）/ C2 定制业务赋能（先萃取再构建，必含能力模型与水平；<code>VR-TIER03</code>）。</li>
  <li><strong>非 AVG 50%</strong>：导入 + 复盘结构化产出 + 工作场景应用 + 个人行为改变计划（IDP）<strong>四活动齐全</strong>（<code>VR-TIER02</code> BLOCKER）；应用与 IDP 模板独立、均指向真实场景。</li>
  <li><strong>Pipeline 第 0 阶段（Intake &amp; Tiering）</strong>：输入检查（完整 / 脱敏 / 版权）→ 分层判定（tier/track + 50/50 时间预算 + 定价/工时锚）→ 产出 Intake 卡片 → 才进入 Course → Capability → …（<code>VR-INTAKE01</code> BLOCKER：无判定不得进入 Pipeline）。</li>
  <li><strong>清单嵌入</strong>：<code>Sample_Input_Checklist</code> 第 11 项「分层与前置判定」由课程主人预先自评，Intake Agent 复核后写入 AVG Project JSON 的 <code>kit.intake</code>。</li>
</ul>
</section>
"""

def inject_html():
    s = open(HTML, encoding="utf-8").read()
    if "课程分层 3×2 矩阵（设计要求" in s:
        print("PRD.html 已含分层要求段，跳过"); return
    s = s.replace("</body>", HTML_BLOCK + "\n</body>", 1)
    open(HTML, "w", encoding="utf-8").write(s)
    print("PRD.html 已注入分层要求段")

def set_cjk(run, name="Microsoft YaHei"):
    run.font.name = name
    r = run._element; rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'), {}); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def inject_docx():
    d = Document(DOCX)
    if any("课程分层 3×2 矩阵（设计要求" in p.text for p in d.paragraphs):
        print("PRD.docx 已含分层要求段，跳过"); return
    h = d.add_heading("课程分层 3×2 矩阵（设计要求 · 设计前前置判定）", level=1)
    for r in h.runs: set_cjk(r)
    p1 = d.add_paragraph(
        "设计要求：任何课程进入平台，必须先完成【课程准入与分层判定】（Course Intake & Tiering），"
        "输出 tier/track 判定后，主流程（Pipeline）才启动。该判定是 Pipeline 的第 0 阶段前置输入，不是可选步骤。"
        "完整矩阵与盲区见 Course_Tiered_Matrix.docx。")
    for r in p1.runs: set_cjk(r)
    for t in [
        "Tier（3 层）：L1 轻量切片 3h / L2 标准单元 7h / L3 深度系列 14h；AVG 探索体验严格占总时长 50%（VR-TIER01 BLOCKER）。",
        "Track（2 类）：C1 通用改编（已有课程转译）/ C2 定制业务赋能（先萃取再构建，必含能力模型与水平；VR-TIER03）。",
        "非 AVG 50%：导入 + 复盘结构化产出 + 工作场景应用 + 个人行为改变计划（IDP）四活动齐全（VR-TIER02 BLOCKER）；应用与 IDP 模板独立、均指向真实场景。",
        "Pipeline 第 0 阶段（Intake & Tiering）：输入检查（完整/脱敏/版权）→ 分层判定（tier/track + 50/50 时间预算 + 定价/工时锚）→ 产出 Intake 卡片 → 才进入 Course → Capability → …（VR-INTAKE01 BLOCKER：无判定不得进入 Pipeline）。",
        "清单嵌入：Sample_Input_Checklist 第 11 项「分层与前置判定」由课程主人预先自评，Intake Agent 复核后写入 AVG Project JSON 的 kit.intake。",
    ]:
        p = d.add_paragraph(t, style='List Bullet')
        for r in p.runs: set_cjk(r)
    d.save(DOCX)
    print("PRD.docx 已注入分层要求段")

if __name__ == "__main__":
    inject_html()
    inject_docx()
